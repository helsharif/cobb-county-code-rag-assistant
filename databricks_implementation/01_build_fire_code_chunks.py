# Databricks notebook source
# MAGIC %md
# MAGIC # Setup and Dependencies

# COMMAND ----------

# MAGIC %pip install pypdf python-docx pandas
# MAGIC dbutils.library.restartPython()

# COMMAND ----------

# MAGIC %md
# MAGIC ## Paths

# COMMAND ----------

CATALOG = "cobb_rag"
SCHEMA = "fire_code"

VOLUME_PATH = "/Volumes/cobb_rag/fire_code/source_docs"
CHUNKS_TABLE = "cobb_rag.fire_code.docling_chunks"

# COMMAND ----------

# find all supported files

SUPPORTED_EXTENSIONS = (
    ".pdf",
    ".docx",
    ".doc",
    ".html",
    ".htm",
    ".md",
    ".txt",
)

def list_files_recursive(path):
    files = []
    for item in dbutils.fs.ls(path):
        if item.isDir():
            files.extend(list_files_recursive(item.path))
        elif item.path.lower().endswith(SUPPORTED_EXTENSIONS):
            files.append(item.path)
    return files

files = list_files_recursive(VOLUME_PATH)

display(spark.createDataFrame([(f,) for f in files], ["path"]))

# COMMAND ----------

# MAGIC %md
# MAGIC ## Confirm files found

# COMMAND ----------

# Show how many supported files were found in the volume.
# This confirms that recursive folder discovery is working before extraction starts.

print(f"Found {len(files)} supported files")

for path in files[:20]:
    print(path)

# COMMAND ----------

# MAGIC %md
# MAGIC ## Create Helpers For File Paths And Text Cleaning

# COMMAND ----------

import os
import re
import uuid
from datetime import datetime, UTC

from pypdf import PdfReader
from docx import Document

# Databricks file listings may return paths like:
# dbfs:/Volumes/cobb_rag/fire_code/source_docs/...
#
# Python file readers usually need:
# /Volumes/cobb_rag/fire_code/source_docs/...
#
# This helper converts Databricks paths to local filesystem paths.
def to_local_path(path):
    if path.startswith("dbfs:"):
        return path.replace("dbfs:", "", 1)
    return path


# Normalize whitespace while preserving paragraph breaks where possible.
# This makes chunks cleaner for vector search and keyword search.
def clean_text(text):
    text = text or ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


# Get the top-level source folder, such as:
# cobb_county_fire, cobb_municode, applicable_codes, etc.
def source_folder_for(path):
    clean_path = path.replace("dbfs:", "", 1)
    relative = clean_path.replace(VOLUME_PATH + "/", "")
    return relative.split("/")[0]


# Create a stable document id from the filename.
# Example: "Fire Code.pdf" -> "fire_code_pdf"
def doc_id_for(path):
    filename = os.path.basename(path)
    return re.sub(r"[^a-zA-Z0-9_]+", "_", filename).strip("_").lower()

# COMMAND ----------

# MAGIC %md
# MAGIC ## Add File Readers

# COMMAND ----------

# Read a PDF one page at a time.
# Returning page-level text lets us preserve page numbers for citations later.
def read_pdf(path):
    local_path = to_local_path(path)
    reader = PdfReader(local_path)

    pages = []

    for page_number, page in enumerate(reader.pages, start=1):
        text = clean_text(page.extract_text())

        if text:
            pages.append({
                "page_start": page_number,
                "page_end": page_number,
                "text": text,
            })

    return pages


# Read DOCX files.
# python-docx does not support old .doc files, only .docx.
def read_docx(path):
    local_path = to_local_path(path)
    document = Document(local_path)

    paragraphs = []
    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            paragraphs.append(text)

    combined = clean_text("\n\n".join(paragraphs))

    if not combined:
        return []

    return [{
        "page_start": None,
        "page_end": None,
        "text": combined,
    }]


# Read plain text, Markdown, and simple HTML files.
# This is intentionally lightweight for the first working pipeline.
def read_text_like_file(path):
    local_path = to_local_path(path)

    with open(local_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()

    # Remove simple HTML tags if the file is HTML.
    if path.lower().endswith((".html", ".htm")):
        text = re.sub(r"<script.*?</script>", " ", text, flags=re.I | re.S)
        text = re.sub(r"<style.*?</style>", " ", text, flags=re.I | re.S)
        text = re.sub(r"<[^>]+>", " ", text)

    text = clean_text(text)

    if not text:
        return []

    return [{
        "page_start": None,
        "page_end": None,
        "text": text,
    }]

# COMMAND ----------

# MAGIC %md
# MAGIC ## Add Chunking

# COMMAND ----------

# Split text into overlapping chunks.
# Smaller chunks improve retrieval precision; overlap prevents losing context at boundaries.
def chunk_text(text, max_chars=1800, overlap=250):
    text = clean_text(text)
    chunks = []

    start = 0

    while start < len(text):
        end = min(start + max_chars, len(text))

        # Prefer to end chunks at paragraph boundaries when possible.
        boundary = text.rfind("\n\n", start, end)
        if boundary > start + 600:
            end = boundary

        chunk = text[start:end].strip()

        if chunk:
            chunks.append(chunk)

        if end >= len(text):
            break

        start = max(0, end - overlap)

    return chunks

# COMMAND ----------

# MAGIC %md
# MAGIC ## Test One File Before Processing Everything, if it works, continue

# COMMAND ----------

# Test extraction on one file before creating the full Delta table.
# This catches path or parser issues early.

test_path = files[0]
print(f"Testing: {test_path}")

lower = test_path.lower()

if lower.endswith(".pdf"):
    extracted = read_pdf(test_path)
elif lower.endswith(".docx"):
    extracted = read_docx(test_path)
elif lower.endswith((".txt", ".md", ".html", ".htm")):
    extracted = read_text_like_file(test_path)
else:
    extracted = []

print(f"Extracted sections/pages: {len(extracted)}")

if extracted:
    print(extracted[0]["text"][:2000])

# COMMAND ----------

# MAGIC %md
# MAGIC # Create Or Reset The Delta Table
# MAGIC Run this only when you intentionally want a fresh table.

# COMMAND ----------

from pyspark.sql.types import (
    StructType,
    StructField,
    StringType,
    IntegerType,
    TimestampType,
)

# Use a non-Docling table name for this pypdf/python-docx pipeline.
CHUNKS_TABLE = "cobb_rag.fire_code.document_chunks"
FAILURES_TABLE = "cobb_rag.fire_code.document_ingestion_failures"

# Create catalog and schema if needed.
spark.sql(f"CREATE CATALOG IF NOT EXISTS {CATALOG}")
spark.sql(f"CREATE SCHEMA IF NOT EXISTS {CATALOG}.{SCHEMA}")

# Define the table schema explicitly so appends work even when page numbers are null.
chunk_schema = StructType([
    StructField("id", StringType(), False),
    StructField("doc_id", StringType(), True),
    StructField("source_folder", StringType(), True),
    StructField("source_path", StringType(), True),
    StructField("source_file", StringType(), True),
    StructField("page_start", IntegerType(), True),
    StructField("page_end", IntegerType(), True),
    StructField("chunk_index", IntegerType(), True),
    StructField("text", StringType(), True),
    StructField("parser", StringType(), True),
    StructField("created_at", TimestampType(), True),
])

empty_df = spark.createDataFrame([], chunk_schema)

(
    empty_df.write
    .format("delta")
    .mode("overwrite")
    .option("overwriteSchema", "true")
    .saveAsTable(CHUNKS_TABLE)
)

spark.sql(f"""
ALTER TABLE {CHUNKS_TABLE}
SET TBLPROPERTIES (delta.enableChangeDataFeed = true)
""")

print(f"Created fresh table: {CHUNKS_TABLE}")

# COMMAND ----------

# MAGIC %md
# MAGIC # Chunking

# COMMAND ----------

# MAGIC %md
# MAGIC ## Process one subfolder first
# MAGIC Start with one folder so the run is easy to debug.
# MAGIC
# MAGIC After this works, repeat with the other folders or use all files.
# MAGIC

# COMMAND ----------

folder_to_process = "/Volumes/cobb_rag/fire_code/source_docs/cobb_county_fire"

batch_files = list_files_recursive(folder_to_process)

print(f"Files in batch: {len(batch_files)}")

for path in batch_files:
    print(path)

# COMMAND ----------

failures = []

for path in batch_files:
    try:
        print(f"Processing: {path}")

        lower = path.lower()

        if lower.endswith(".pdf"):
            extracted_items = read_pdf(path)
            parser_name = "pypdf"
        elif lower.endswith(".docx"):
            extracted_items = read_docx(path)
            parser_name = "python-docx"
        elif lower.endswith((".txt", ".md", ".html", ".htm")):
            extracted_items = read_text_like_file(path)
            parser_name = "text"
        else:
            print(f"Skipping unsupported file type: {path}")
            continue

        filename = os.path.basename(path)
        folder = source_folder_for(path)
        doc_id = doc_id_for(path)

        rows = []
        chunk_index = 0

        for item in extracted_items:
            page_chunks = chunk_text(item["text"])

            for chunk in page_chunks:
                rows.append({
                    "id": str(uuid.uuid4()),
                    "doc_id": doc_id,
                    "source_folder": folder,
                    "source_path": path,
                    "source_file": filename,
                    "page_start": item["page_start"],
                    "page_end": item["page_end"],
                    "chunk_index": chunk_index,
                    "text": chunk,
                    "parser": parser_name,
                    "created_at": datetime.now(UTC),
                })

                chunk_index += 1

        if rows:
            df = spark.createDataFrame(rows, schema=chunk_schema)

            (
                df.write
                .format("delta")
                .mode("append")
                .saveAsTable(CHUNKS_TABLE)
            )

        print(f"Saved {len(rows)} chunks from {filename}")

    except Exception as e:
        failures.append({
            "source_path": path,
            "error": str(e),
            "created_at": datetime.now(UTC),
        })

        print(f"FAILED: {path}")
        print(e)

# COMMAND ----------

# MAGIC %md
# MAGIC ### Inspect results of test run

# COMMAND ----------


spark.sql(f"""
SELECT
  source_folder,
  COUNT(*) AS chunk_count,
  COUNT(DISTINCT source_file) AS file_count
FROM {CHUNKS_TABLE}
GROUP BY source_folder
ORDER BY source_folder
""").show(truncate=False)

# COMMAND ----------

# MAGIC %md
# MAGIC ### preview chunks:

# COMMAND ----------

display(
    spark.table(CHUNKS_TABLE)
    .select(
        "source_file",
        "page_start",
        "page_end",
        "chunk_index",
        "text",
    )
    .limit(20)
)

# COMMAND ----------

# MAGIC %md
# MAGIC ## Process all folders
# MAGIC If the one folder test worked, reset the delta table (see previous section of notebook) and now process all folders

# COMMAND ----------

# Imports used for extraction, chunking, metadata, and schema definitions.

import os
import re
import uuid
from datetime import datetime, UTC

from pypdf import PdfReader
from docx import Document

from pyspark.sql.types import (
    StructType,
    StructField,
    StringType,
    IntegerType,
    TimestampType,
)

CATALOG = "cobb_rag"
SCHEMA = "fire_code"

VOLUME_PATH = "/Volumes/cobb_rag/fire_code/source_docs"

CHUNKS_TABLE = "cobb_rag.fire_code.document_chunks"
FAILURES_TABLE = "cobb_rag.fire_code.document_ingestion_failures"

SUPPORTED_EXTENSIONS = (
    ".pdf",
    ".docx",
    ".html",
    ".htm",
    ".md",
    ".txt",
)

# COMMAND ----------

# MAGIC %md
# MAGIC ### Helpers

# COMMAND ----------

# Recursively find supported files in a Databricks Volume.
def list_files_recursive(path):
    files = []

    for item in dbutils.fs.ls(path):
        if item.isDir():
            files.extend(list_files_recursive(item.path))
        elif item.path.lower().endswith(SUPPORTED_EXTENSIONS):
            files.append(item.path)

    return files


# Convert Databricks dbfs:/ paths into local paths readable by Python libraries.
def to_local_path(path):
    if path.startswith("dbfs:"):
        return path.replace("dbfs:", "", 1)
    return path


# Normalize whitespace while preserving paragraph boundaries.
def clean_text(text):
    text = text or ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


# Extract the first folder below the source volume as a source category.
def source_folder_for(path):
    clean_path = path.replace("dbfs:", "", 1)
    relative = clean_path.replace(VOLUME_PATH + "/", "")
    return relative.split("/")[0]


# Create a stable document id from the source filename.
def doc_id_for(path):
    filename = os.path.basename(path)
    return re.sub(r"[^a-zA-Z0-9_]+", "_", filename).strip("_").lower()

# COMMAND ----------

# MAGIC %md
# MAGIC ### Readers And Chunker

# COMMAND ----------

# Extract PDF text page by page so citations can keep page numbers.
def read_pdf(path):
    local_path = to_local_path(path)
    reader = PdfReader(local_path)

    pages = []

    for page_number, page in enumerate(reader.pages, start=1):
        text = clean_text(page.extract_text())

        if text:
            pages.append({
                "page_start": page_number,
                "page_end": page_number,
                "text": text,
            })

    return pages


# Extract paragraph text from DOCX files.
def read_docx(path):
    local_path = to_local_path(path)
    document = Document(local_path)

    paragraphs = []

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            paragraphs.append(text)

    combined = clean_text("\n\n".join(paragraphs))

    if not combined:
        return []

    return [{
        "page_start": None,
        "page_end": None,
        "text": combined,
    }]


# Extract text from TXT, Markdown, and simple HTML files.
def read_text_like_file(path):
    local_path = to_local_path(path)

    with open(local_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()

    if path.lower().endswith((".html", ".htm")):
        text = re.sub(r"<script.*?</script>", " ", text, flags=re.I | re.S)
        text = re.sub(r"<style.*?</style>", " ", text, flags=re.I | re.S)
        text = re.sub(r"<[^>]+>", " ", text)

    text = clean_text(text)

    if not text:
        return []

    return [{
        "page_start": None,
        "page_end": None,
        "text": text,
    }]


# Split extracted text into overlapping chunks for retrieval.
def chunk_text(text, max_chars=1800, overlap=250):
    text = clean_text(text)
    chunks = []
    start = 0

    while start < len(text):
        end = min(start + max_chars, len(text))

        # Prefer paragraph boundaries when they are not too close to the start.
        boundary = text.rfind("\n\n", start, end)
        if boundary > start + 600:
            end = boundary

        chunk = text[start:end].strip()

        if chunk:
            chunks.append(chunk)

        if end >= len(text):
            break

        start = max(0, end - overlap)

    return chunks

# COMMAND ----------

# MAGIC %md
# MAGIC ### Schemas

# COMMAND ----------

# Explicit schemas keep Spark appends stable when optional values like page numbers are null.

chunk_schema = StructType([
    StructField("id", StringType(), False),
    StructField("doc_id", StringType(), True),
    StructField("source_folder", StringType(), True),
    StructField("source_path", StringType(), True),
    StructField("source_file", StringType(), True),
    StructField("page_start", IntegerType(), True),
    StructField("page_end", IntegerType(), True),
    StructField("chunk_index", IntegerType(), True),
    StructField("text", StringType(), True),
    StructField("parser", StringType(), True),
    StructField("created_at", TimestampType(), True),
])

failure_schema = StructType([
    StructField("source_path", StringType(), True),
    StructField("error", StringType(), True),
    StructField("created_at", TimestampType(), True),
])

# COMMAND ----------

# MAGIC %md
# MAGIC ### Reset Tables
# MAGIC

# COMMAND ----------

# Reset the output tables for a clean full ingestion run.
# Only run this cell when you intentionally want to rebuild from scratch.

spark.sql(f"CREATE CATALOG IF NOT EXISTS {CATALOG}")
spark.sql(f"CREATE SCHEMA IF NOT EXISTS {CATALOG}.{SCHEMA}")

empty_chunks_df = spark.createDataFrame([], chunk_schema)
empty_failures_df = spark.createDataFrame([], failure_schema)

(
    empty_chunks_df.write
    .format("delta")
    .mode("overwrite")
    .option("overwriteSchema", "true")
    .saveAsTable(CHUNKS_TABLE)
)

(
    empty_failures_df.write
    .format("delta")
    .mode("overwrite")
    .option("overwriteSchema", "true")
    .saveAsTable(FAILURES_TABLE)
)

spark.sql(f"""
ALTER TABLE {CHUNKS_TABLE}
SET TBLPROPERTIES (delta.enableChangeDataFeed = true)
""")

print(f"Reset table: {CHUNKS_TABLE}")
print(f"Reset table: {FAILURES_TABLE}")

# COMMAND ----------

# MAGIC %md
# MAGIC ### Find All Files

# COMMAND ----------

# Find all source files across every subfolder in the volume.

all_files = list_files_recursive(VOLUME_PATH)

print(f"Found {len(all_files)} supported files")

display(
    spark.createDataFrame([(path,) for path in all_files], ["source_path"])
)

# COMMAND ----------

# MAGIC %md
# MAGIC ### Full Ingestion Loop

# COMMAND ----------

# Process every supported file and append chunks after each file.
# Appending file-by-file preserves progress if one later file fails.

total_chunks = 0
total_failures = 0

for file_number, path in enumerate(all_files, start=1):
    try:
        print(f"[{file_number}/{len(all_files)}] Processing: {path}")

        lower = path.lower()

        if lower.endswith(".pdf"):
            extracted_items = read_pdf(path)
            parser_name = "pypdf"
        elif lower.endswith(".docx"):
            extracted_items = read_docx(path)
            parser_name = "python-docx"
        elif lower.endswith((".txt", ".md", ".html", ".htm")):
            extracted_items = read_text_like_file(path)
            parser_name = "text"
        else:
            print(f"Skipping unsupported file type: {path}")
            continue

        filename = os.path.basename(path)
        folder = source_folder_for(path)
        doc_id = doc_id_for(path)

        rows = []
        chunk_index = 0

        for item in extracted_items:
            chunks = chunk_text(item["text"])

            for chunk in chunks:
                rows.append({
                    "id": str(uuid.uuid4()),
                    "doc_id": doc_id,
                    "source_folder": folder,
                    "source_path": path,
                    "source_file": filename,
                    "page_start": item["page_start"],
                    "page_end": item["page_end"],
                    "chunk_index": chunk_index,
                    "text": chunk,
                    "parser": parser_name,
                    "created_at": datetime.now(UTC),
                })

                chunk_index += 1

        if rows:
            chunks_df = spark.createDataFrame(rows, schema=chunk_schema)

            (
                chunks_df.write
                .format("delta")
                .mode("append")
                .saveAsTable(CHUNKS_TABLE)
            )

        total_chunks += len(rows)

        print(f"Saved {len(rows)} chunks from {filename}")

    except Exception as e:
        total_failures += 1

        failure_rows = [{
            "source_path": path,
            "error": str(e),
            "created_at": datetime.now(UTC),
        }]

        failures_df = spark.createDataFrame(failure_rows, schema=failure_schema)

        (
            failures_df.write
            .format("delta")
            .mode("append")
            .saveAsTable(FAILURES_TABLE)
        )

        print(f"FAILED: {path}")
        print(e)

print(f"Finished ingestion")
print(f"Total chunks saved: {total_chunks}")
print(f"Total file failures: {total_failures}")

# COMMAND ----------

# MAGIC %md
# MAGIC ### Inspect Results By Folder

# COMMAND ----------

# Summarize the final chunk table by source folder.

spark.sql(f"""
SELECT
  source_folder,
  COUNT(*) AS chunk_count,
  COUNT(DISTINCT source_file) AS file_count
FROM {CHUNKS_TABLE}
GROUP BY source_folder
ORDER BY source_folder
""").show(truncate=False)

# COMMAND ----------

# MAGIC %md
# MAGIC ### Inspect Sample Chunks

# COMMAND ----------

# Preview chunks to make sure extraction quality looks reasonable.

display(
    spark.table(CHUNKS_TABLE)
    .select(
        "source_folder",
        "source_file",
        "page_start",
        "page_end",
        "chunk_index",
        "text",
    )
    .orderBy("source_folder", "source_file", "chunk_index")
    .limit(30)
)

# COMMAND ----------

# Review any files that failed extraction.

failure_count = spark.table(FAILURES_TABLE).count()

print(f"Failure count: {failure_count}")

if failure_count > 0:
    display(spark.table(FAILURES_TABLE))

# COMMAND ----------

# MAGIC %md
# MAGIC ### Final Count

# COMMAND ----------

# Final table size check before creating the AI Search hybrid index.

spark.sql(f"""
SELECT COUNT(*) AS total_chunks
FROM {CHUNKS_TABLE}
""").show()

# COMMAND ----------

# MAGIC %md
# MAGIC After this finishes, your table for the AI Search index will be:
# MAGIC
# MAGIC ```
# MAGIC cobb_rag.fire_code.document_chunks
# MAGIC ```
# MAGIC
# MAGIC When you create the index, use:
# MAGIC
# MAGIC ```
# MAGIC Primary key: id
# MAGIC Text / embedding source column: text
# MAGIC Index type: Hybrid
# MAGIC Sync mode: Triggered
# MAGIC ```
# MAGIC
# MAGIC Sync these columns:
# MAGIC
# MAGIC ```
# MAGIC id
# MAGIC text
# MAGIC doc_id
# MAGIC source_folder
# MAGIC source_path
# MAGIC source_file
# MAGIC page_start
# MAGIC page_end
# MAGIC chunk_index
# MAGIC parser
# MAGIC created_at
# MAGIC ```

# COMMAND ----------

# MAGIC %md
# MAGIC # Databricks AI Search Hybrid Index
# MAGIC
# MAGIC Create the Databricks AI Search Hybrid index from your chunk table, then smoke-test retrieval.

# COMMAND ----------

# MAGIC %md
# MAGIC ## Verify Table Is Ready

# COMMAND ----------

# Confirm the chunk table exists and has rows before creating the search index.

CHUNKS_TABLE = "cobb_rag.fire_code.document_chunks"

spark.sql(f"""
SELECT COUNT(*) AS total_chunks
FROM {CHUNKS_TABLE}
""").show()

# COMMAND ----------

# MAGIC %md
# MAGIC ## Check that important metadata columns are populated.
# MAGIC

# COMMAND ----------

# Check that important metadata columns are populated.

spark.sql(f"""
SELECT
  source_folder,
  COUNT(*) AS chunk_count,
  COUNT(DISTINCT source_file) AS file_count
FROM {CHUNKS_TABLE}
GROUP BY source_folder
ORDER BY source_folder
""").show(truncate=False)

# COMMAND ----------

# MAGIC %md
# MAGIC # Create An AI Search Endpoint
# MAGIC
# MAGIC In the Databricks UI:
# MAGIC - Go to Compute.
# MAGIC - Click the AI Search tab.
# MAGIC - Click Create.
# MAGIC - Name it something like:
# MAGIC ```cobb-fire-search-endpoint```
# MAGIC - Choose Standard if prompted.
# MAGIC - Create it and wait until it is online.
# MAGIC
# MAGIC
# MAGIC Databricks docs say Standard is the usual choice for smaller indexes and lower latency. Storage-optimized is more for huge indexes.